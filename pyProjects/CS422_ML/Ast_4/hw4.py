import pandas as pd
from sklearn import datasets, svm
from sklearn.model_selection import cross_val_score
import numpy as np
from docx import Document

# Load the CSV data
mnist_data = pd.read_csv('MNIST_HW4.csv')

# Prepare data
feature_data = mnist_data.drop('label', axis=1)
target_data = mnist_data['label']

# Initialize SVM classifiers with different kernels
kernel_types = ['linear', 'poly', 'rbf']
svm_classifiers = [svm.SVC(kernel=k, gamma='scale') for k in kernel_types]

# Perform 5-fold cross-validation and compute accuracy
for svm_clf, kernel_type in zip(svm_classifiers, kernel_types):
    cv_scores = cross_val_score(svm_clf, feature_data, target_data, cv=5)
    print(f"Accuracy for {kernel_type} kernel: {np.mean(cv_scores)}")

# Create a new Word doc
hw4_doc = Document()

# Add the content to the document
hw4_doc.add_paragraph('Brandon Timok')
hw4_doc.add_paragraph('CS 422 Machine Learning')
hw4_doc.add_paragraph('Junggab Son')
hw4_doc.add_paragraph('Homework 4').alignment = 1

hw4_doc.add_paragraph('Used libraries and their purposes:')
hw4_doc.add_paragraph('sklearn: Used for implementing SVM with multiple kernels and computing accuracy using 5-fold CV.')


hw4_doc.add_paragraph('Screenshots:')
hw4_doc.add_picture('hw4shot.png')

hw4_doc.add_paragraph('What I learned from this assignment:')
hw4_doc.add_paragraph('I learned how to use SVMs with different kernels (linear, poly, rbf) for multi-class classification and computing accuracy.')

# Save document
hw4_doc.save('hw4.docx')