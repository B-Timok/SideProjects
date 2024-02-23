import numpy as np
import pandas as pd
from collections import Counter
from docx import Document

# Step 1: Load the training and test data
train_data = pd.read_csv("MNIST_training.csv")
test_data = pd.read_csv("MNIST_test.csv")

# Define a function to calculate Euclidean distance
def euclidean_distance(a, b):
    return np.sqrt(np.sum((a - b) ** 2, axis=1))

# Step 3: Compute distances/similarities between test and training data
distances = np.zeros((test_data.shape[0], train_data.shape[0]))
for i in range(test_data.shape[0]):
    distances[i] = euclidean_distance(test_data.values[i, 1:], train_data.values[:, 1:])

# Step 5: Compare prediction with ground truth
ground_truth = test_data.values[:, 0]

# Initialize an empty DataFrame to store the results
results = pd.DataFrame(columns=['K', 'Correctly Classified', 'Incorrectly Classified', 'Accuracy'])

# Test specific values of K
for K in [1, 3, 5, 7, 9]:  # Change the list as needed
    # Step 4: Find K-nearest neighbors and decide majority class
    knn_indices = np.argpartition(distances, K, axis=1)[:, :K]
    knn_labels = train_data.values[knn_indices, 0]
    predicted_labels = np.apply_along_axis(lambda x: np.bincount(x).argmax(), axis=1, arr=knn_labels)

    # Compute accuracy
    correctly_classified = np.sum(predicted_labels == ground_truth)
    incorrectly_classified = len(test_data) - correctly_classified
    accuracy = correctly_classified / len(test_data)

    # Add the results to the DataFrame
    results.loc[len(results)] = {'K': K, 'Correctly Classified': correctly_classified, 'Incorrectly Classified': incorrectly_classified, 'Accuracy': accuracy}

# Print the DataFrame
print(results)

# Create a new Word document
doc = Document()

# Add a title
doc.add_heading('K-Nearest Neighbors (KNN) Implementation', level=1)

# Add the detailed explanation
doc.add_heading('Explanation', level=2)
doc.add_paragraph("""
This program implements K-nearest neighbors (KNN) algorithm from scratch using Python. 
It takes two datasets, MNIST_training.csv and MNIST_test.csv, and follows the steps below:
1. Load the training and test data.
2. Calculate distances (Euclidean, Manhattan, or Cosine similarity) between test and training data.
3. Find the K-nearest neighbors and decide the majority class.
4. Compare the prediction with the ground truth in the test data.
5. Compute accuracy by counting correctly and incorrectly classified samples.
6. Repeat the process for different values of K and display the accuracy results.

""")

# Add tables for K values and test accuracy
doc.add_heading('Results', level=2)
doc.add_paragraph("")

# Convert the 'results' DataFrame to a table
results_table = doc.add_table(rows=len(results)+1, cols=4)
header_cells = results_table.rows[0].cells
header_cells[0].text = 'K'
header_cells[1].text = 'Correctly Classified'
header_cells[2].text = 'Incorrectly Classified'
header_cells[3].text = 'Accuracy'

# Add data to the table
for i in range(len(results)):
    row_cells = results_table.rows[i+1].cells
    row_cells[0].text = str(results.loc[i, 'K'])
    row_cells[1].text = str(results.loc[i, 'Correctly Classified'])
    row_cells[2].text = str(results.loc[i, 'Incorrectly Classified'])
    row_cells[3].text = str(results.loc[i, 'Accuracy'])

# Save the document
doc.save('KNN_Results.docx')